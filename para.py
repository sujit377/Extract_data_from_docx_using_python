import docx;
import xlsxwriter;
doc = docx.Document("G:/Free Lancer/3310264_Daily Bid Sheet 2_12.docx");
t = doc.tables
w = xlsxwriter.Workbook("G:/Free Lancer/test.xlsx");
w1 = w.add_worksheet()
w1.write(0,0,"NAME")
w1.write(0,1,"PHONE NO.")
w1.write(0,2,"ADDRESS")
for t1 in t:
		col = 1;
		for row in t1.rows:
			data = []
			cell = row.cells[0]
			l = len(cell.paragraphs)
			z = 0;
			if(l<2 and len(cell.paragraphs[0].text)==0):
				continue;
			else:
				p1 =cell.paragraphs[0]
				s1 = p1.text
				s1 = s1.split(" ");
				s1 = s1[0]
				if(s1.isalpha()):
					data.append(p1.text);
					if(l<3 and len(cell.paragraphs[1].text)==0):
						continue;
					else:
						z=z+1;
				else:
					continue;
				p2 =cell.paragraphs[1]
				s2 = p2.text
				s2 = s2.split(" ");
				s2 = s2[0]
				if(not( s2.isalnum())):
					data.append(p2.text);
					z = 2
				else:
					data.append("")
				m = ""
				while(z<l and len(cell.paragraphs[z].text)>0):
					p =cell.paragraphs[z]
					m =m + p.text + " "
					z+=1;
				data.append(m)
			w1.write(col,0,data[0])
			w1.write(col,1,data[1])
			w1.write(col,2,data[2])
			col+=1;

w.close()
		
			

			

				
				
				